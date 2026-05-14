"""Parse a .docx Support Notes content template into structured slot data.

The document may be structured as paragraphs and/or tables. All text and
hyperlinks are extracted in document order, then passed to Claude Haiku
for slot mapping — the same model and prompt used by the .eml path.

Public entry point
------------------
``parse(source) -> (SupportNotesData, List[EmlAttachment], Optional[str])``

Like pptx_parser, product / language / year / month are left at their
defaults; the caller sets language from the UI picker.
"""

from __future__ import annotations

import io
import json
import os
import re
from pathlib import Path
from typing import IO, List, Optional, Tuple, Union

from builder.models import EmlAttachment, SupportNotesData

PROMPT_PATH = Path(__file__).resolve().parent / "prompts" / "extract_slots.txt"
_MODEL_ID = "claude-haiku-4-5-20251001"

try:
    import truststore
    truststore.inject_into_ssl()
except ImportError:
    pass


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def parse(
    source: Union[str, Path, bytes, IO[bytes]],
) -> Tuple[SupportNotesData, List[EmlAttachment], Optional[str]]:
    """Parse a .docx file and return ``(data, headshot_images, ai_error)``.

    ``data`` has content slots populated by Claude Haiku where possible.
    ``headshot_images`` contains embedded images (≥50 KB) as EmlAttachment objects.
    ``ai_error`` is None on success or a human-readable warning when AI failed.
    """
    doc = _load_docx(source)
    body_text = _extract_body_text(doc)
    images = _collect_images(doc)

    ai_error: Optional[str] = None
    if os.environ.get("ANTHROPIC_API_KEY", "").strip():
        try:
            data = _ai_extract(body_text)
        except Exception as exc:
            ai_error = _friendly_error(exc)
            data = SupportNotesData()
    else:
        data = SupportNotesData()

    return data, images, ai_error


# ---------------------------------------------------------------------------
# Document loading
# ---------------------------------------------------------------------------

def _load_docx(source: Union[str, Path, bytes, IO[bytes]]):
    from docx import Document
    if isinstance(source, (str, Path)):
        return Document(str(source))
    if isinstance(source, bytes):
        return Document(io.BytesIO(source))
    if hasattr(source, "read"):
        return Document(source)
    raise TypeError(f"Unsupported source type for .docx parse: {type(source)!r}")


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _rel_map(part) -> dict:
    return {
        r_id: rel.target_ref
        for r_id, rel in part.rels.items()
        if hasattr(rel, "target_ref")
    }


def _para_text(para, rels: dict) -> str:
    """Return paragraph text with hyperlinks rendered as 'text [URL]'."""
    from docx.oxml.ns import qn
    parts: list[str] = []
    for child in para._element:
        tag = child.tag
        if tag == qn("w:r"):
            t = child.find(qn("w:t"))
            if t is not None and t.text:
                parts.append(t.text)
        elif tag == qn("w:hyperlink"):
            r_id = child.get(qn("r:id"))
            url = rels.get(r_id, "") if r_id else ""
            link_parts: list[str] = []
            for r in child.iter(qn("w:r")):
                t = r.find(qn("w:t"))
                if t is not None and t.text:
                    link_parts.append(t.text)
            link_text = "".join(link_parts)
            if link_text:
                parts.append(f"{link_text} [{url}]" if url else link_text)
    return "".join(parts).strip()


def _extract_body_text(doc) -> str:
    """Walk the document body in order, converting paragraphs and tables to text."""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    rels = _rel_map(doc.part)
    sections: list[str] = []

    for child in doc.element.body:
        tag = child.tag
        if tag == qn("w:p"):
            para = Paragraph(child, doc)
            text = _para_text(para, rels)
            if text:
                sections.append(text)
        elif tag == qn("w:tbl"):
            tbl = Table(child, doc)
            lines: list[str] = []
            seen_cells: set[int] = set()
            for r_idx, row in enumerate(tbl.rows):
                for c_idx, cell in enumerate(row.cells):
                    # Skip repeated cells caused by merged-cell spanning
                    cell_id = id(cell._tc)
                    if cell_id in seen_cells:
                        continue
                    seen_cells.add(cell_id)
                    cell_lines: list[str] = []
                    for para in cell.paragraphs:
                        line = _para_text(para, rels)
                        if line:
                            cell_lines.append(line)
                    cell_text = "\n".join(cell_lines).strip()
                    if cell_text:
                        lines.append(f"  [{r_idx},{c_idx}]: {cell_text}")
            if lines:
                sections.append("TABLE:\n" + "\n".join(lines))

    return "\n\n".join(sections)


# ---------------------------------------------------------------------------
# AI extraction
# ---------------------------------------------------------------------------

def _ai_extract(body_text: str) -> SupportNotesData:
    from anthropic import Anthropic
    from builder.ai_extractor import _data_from_ai_dict

    client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    system_prompt = PROMPT_PATH.read_text(encoding="utf-8")

    user_msg = (
        "EDITOR_NOTES (plain text extracted from a .docx content template; "
        "hyperlinks are annotated as 'link text [URL]'; "
        "tables are shown as [row,col] cell grid):\n"
        + body_text
        + "\n\n=========================\n\n"
        + "EMAIL_PREVIEW (HTML — the mock-up of the final email):\n"
        + "(none — this is a .docx content template; "
        "extract ALL content from EDITOR_NOTES instead)"
    )

    msg = client.messages.create(
        model=_MODEL_ID,
        max_tokens=8192,
        system=system_prompt,
        messages=[{"role": "user", "content": user_msg}],
    )
    raw = "".join(b.text for b in msg.content if getattr(b, "type", "") == "text").strip()

    if raw.startswith("```"):
        raw = re.sub(r"^```[a-zA-Z]*\s*", "", raw)
        raw = re.sub(r"\s*```\s*$", "", raw)

    try:
        d = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise ValueError(
            f"AI extractor returned non-JSON: {exc}\nFirst 400 chars:\n{raw[:400]}"
        ) from exc

    return _data_from_ai_dict(d)


def _friendly_error(exc: Exception) -> str:
    msg = str(exc)
    low = msg.lower()
    if "ssl" in low or "certificate" in low or "verify" in low:
        return (
            "AI extraction failed: SSL certificate error (corporate proxy). "
            f"Fill in fields manually. (Detail: {msg})"
        )
    if "connection" in low or "timeout" in low or "network" in low:
        return (
            "AI extraction failed: could not reach the Anthropic API. "
            f"Check your network connection. (Detail: {msg})"
        )
    if "json" in low:
        return f"AI extraction failed: model returned invalid JSON. (Detail: {msg})"
    return f"AI extraction failed: {msg}"


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def _collect_images(doc) -> List[EmlAttachment]:
    """Collect embedded images (≥50 KB) from the document's relationships."""
    images: List[EmlAttachment] = []
    seen: set[str] = set()
    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            tp = rel.target_part
            blob = tp.blob
        except Exception:
            continue
        if len(blob) < 50_000:
            continue
        part_name = str(tp.partname)
        if part_name in seen:
            continue
        seen.add(part_name)
        ct = getattr(tp, "content_type", "image/png")
        filename = part_name.rsplit("/", 1)[-1]
        images.append(EmlAttachment(
            filename=filename,
            content_type=ct,
            content_id=None,
            bytes=blob,
        ))
    return images
